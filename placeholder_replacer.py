import os
import re
from docx import Document

class WordPlaceholderReplacer:
    def __init__(self):
        """
        Initializes the WordPlaceholderReplacer with input and output file paths.
        """
        self._template_file = None
        self._output_file = None
        self._document = None
    
    @property
    def save_file(self):
        return self._output_file
    
    @save_file.setter
    def save_file(self, path):
        self._output_file = path

    @property
    def template_file(self):
        return self._template_file
    
    @property
    def document(self):
        return self._document

    @template_file.setter
    def template_file(self, path):
        self._template_file = path
        self._document = Document(path)

    def replace_placeholders(self, replacements, list_delimiter="\n"):
        """
        Replaces placeholders enclosed in {{ and }} throughout the Word document.

        Args:
            replacements (dict): A dictionary where keys are placeholder names without curly braces.
            list_delimiter (str): Delimiter to join list elements if a replacement value is a list.
        """
        # Convert lists in replacements to joined strings using the specified delimiter
        processed_replacements = {
            key: (list_delimiter.join(str(item) for item in value) if isinstance(value, list) else value)
            for key, value in replacements.items()
        }

        # Replace placeholders
        # Paragraphs
        self._process_paragraphs(self.document.paragraphs, processed_replacements)

        # table
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._process_paragraphs(cell.paragraphs, processed_replacements)

        # footer
        for section in self.document.sections:
            self._process_paragraphs(section.footer.paragraphs, processed_replacements)

    def _process_paragraphs(self, paragraphs, replacements):
        """
        Iterates over each paragraph and replaces placeholders.

        Args:
            paragraphs (list): List of docx.paragraph.Paragraph objects.
            replacements (dict): Dictionary with replacement values.
        """
        for paragraph in paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements):
        """
        Replaces placeholders within a paragraph, preserving formatting by processing runs
        and handling split placeholders across runs.
        """
        full_text = "".join(run.text for run in paragraph.runs)
        pattern = re.compile(r'\{\{(.*?)\}\}')

        def replace_placeholder(match):
            key = match.group(1).strip()
            return str(replacements.get(key, match.group(0)))  # Use original if no replacement found

        # Replace placeholders in the entire combined paragraph text
        new_text = pattern.sub(replace_placeholder, full_text)

        # Clear all existing runs and re-apply new text with formatting from the first run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            paragraph.clear()  # Clear the paragraph to prevent duplication
            new_run = paragraph.add_run(new_text)

            # Copy font style from the first run
            if first_run.font:
                new_run.font.bold = first_run.font.bold
                new_run.font.italic = first_run.font.italic
                new_run.font.underline = first_run.font.underline
                new_run.font.name = first_run.font.name
                new_run.font.size = first_run.font.size

    def save(self, filename):
        """Saves the modified Word document to the output file."""
        path = os.path.join(self.save_file, f"{filename}.docx")
        self.document.save(path)


if __name__ == "__main__":
    data = {
    "client_name": "KASSEL RESIDENCES PARAÑAQUE",
    "location": "E. Rodriguez Avenue, La Huerta, Parañaque City",
    "case_number": "NOV-EMB-NCR-2022-1664",
    "date_of_inspection": "31 March 2022",
    "violations": [
        "For having/operating air pollution source Section 1, Rule XIX of the Implementing installments (Diesel Fuel Fire Pump) without Rules and Regulations (IRR) of the corresponding valid 'Permit to Operate' (PO). 'Philippine Clean Air Act of 1999' (RA 8749), as amended by DAO 2004-261",
        "For failure to put a billboard with a message Condition Number 3 of its ECC, pursuant to indicated in the said ECC Condition. the Revised Procedural Manual of DAO No. 30 series of 2003, the Implementing Rules and Regulations of PD 1586.2",
        "Condition Number 4 of its ECC, pursuant to For failure to conform with the provisions of RA 9275, RA 8749, and RA 6969 and its Revised Procedural Manual of DAO No. 30 series of 2003, the Implementing Rules and Regulations of PD 1586.3",
        "For failure to appoint/designate an accredited Condition Number 5 of its ECC, pursuant to Pollution Control Officer that monitors the Revised Procedural Manual of DAO No. 30 series of 2003, the Implementing Rules and Regulations of PD 1586.4",
        "Condition Number 5.2 of its ECC, pursuant to For failure to submit Compliance Monitoring Reports semi-annually every year since the year of 2009.",
        "For operating a facility (wastewater treatment plant) that discharges regulated water pollutants without Discharge Permit. Section 14 of DAO 2005-10 provides that the Department shall require owners or operators of facilities that discharge regulated effluents pursuant to this Act to secure a permit to discharge.",
        "For failure to register as hazardous waste generator considering the Respondent is ordered under Chapter 3.3 of the DENR Administrative Order (DAO) 2013-22, Revised Procedure and Standards for the Management of Hazardous Wastes."
    ],
    "case_number_only": "2022-1664"
    }
    replacer = WordPlaceholderReplacer()
    replacer.template_file = "C:\\Users\\omarg\\OrDraft\\templates\\order_hw_reply.docx"
    replacer.save_file = "."
    replacer.replace_placeholders(data)
    replacer.save()